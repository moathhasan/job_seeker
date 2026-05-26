import re
import io
from typing import List, Dict, Any, Optional
from pypdf import PdfReader
import docx

# A rich set of common skills for matching
COMMON_SKILLS = [
    # Programming Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "golang", "rust", "ruby", 
    "php", "swift", "kotlin", "scala", "r", "sql", "html", "css", "bash", "shell",
    # Frameworks & Libraries
    "react", "angular", "vue", "next.js", "nuxt", "svelte", "express", "node.js", "nodejs",
    "django", "flask", "fastapi", "spring boot", "laravel", "asp.net", "rails", "jquery", 
    "bootstrap", "tailwind", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch",
    # Databases & Cloud
    "mongodb", "postgresql", "mysql", "redis", "sqlite", "dynamodb", "oracle", "firebase",
    "aws", "amazon web services", "azure", "gcp", "google cloud", "docker", "kubernetes", 
    "jenkins", "git", "github", "gitlab", "terraform", "ansible", "ci/cd",
    # Core CS & Methodologies
    "agile", "scrum", "kanban", "oop", "rest api", "graphql", "microservices", "system design",
    "machine learning", "deep learning", "nlp", "computer vision", "data analysis", "data science",
    # Other Roles/Soft Skills
    "project management", "product management", "leadership", "communication", "teamwork", 
    "problem solving", "ui/ux", "figma", "sketch", "adobe xd", "seo", "marketing"
]

class CVParser:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)

    @classmethod
    def parse_cv(cls, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        filename_lower = filename.lower()
        if filename_lower.endswith(".pdf"):
            raw_text = cls.extract_text_from_pdf(file_bytes)
        elif filename_lower.endswith(".docx"):
            raw_text = cls.extract_text_from_docx(file_bytes)
        else:
            # Fallback to plain text
            raw_text = file_bytes.decode("utf-8", errors="ignore")

        # Clean text basic formatting
        clean_text = re.sub(r'\s+', ' ', raw_text)

        # Extract contact info
        email = cls._extract_email(clean_text)
        phone = cls._extract_phone(clean_text)
        name = cls._extract_name(raw_text, filename)
        
        # Extract skills, education, experience
        skills = cls._extract_skills(clean_text)
        education = cls._extract_section(raw_text, ["education", "academic", "university", "college", "degree"])
        experience = cls._extract_section(raw_text, ["experience", "employment", "history", "career", "work history"])

        # Default fallback values for presentation
        if not education:
            education = ["Extracted details from CV: Check raw profile"]
        if not experience:
            experience = ["Extracted details from CV: Check raw profile"]

        return {
            "name": name,
            "email": email,
            "phone": phone,
            "job_title": cls._extract_job_title(raw_text),
            "skills": skills,
            "education": education,
            "experience": experience,
            "raw_text": raw_text
        }

    @staticmethod
    def _extract_email(text: str) -> Optional[str]:
        email_regex = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
        match = re.search(email_regex, text)
        return match.group(0) if match else None

    @staticmethod
    def _extract_phone(text: str) -> Optional[str]:
        # Matches common formats like +1-234-567-8900, (123) 456-7890, 1234567890, etc.
        phone_regex = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        match = re.search(phone_regex, text)
        return match.group(0) if match else None

    @staticmethod
    def _extract_name(raw_text: str, filename: str) -> str:
        # First try to guess from the first few lines of the resume
        lines = [line.strip() for line in raw_text.split("\n") if line.strip()]
        if lines:
            # Check if the first line looks like a name (not too long, doesn't contain contact keywords)
            first_line = lines[0]
            if len(first_line) < 30 and not any(kw in first_line.lower() for kw in ["resume", "cv", "email", "phone", "profile", "contact"]):
                return first_line
        
        # Fallback to filename cleanup
        clean_filename = re.sub(r'(_|cv|resume|\.pdf|\.docx)', ' ', filename, flags=re.IGNORECASE).strip()
        clean_filename = re.sub(r'\s+', ' ', clean_filename)
        return clean_filename.title() if clean_filename else "Candidate Name"

    @staticmethod
    def _extract_skills(text: str) -> List[str]:
        text_lower = text.lower()
        found_skills = []
        for skill in COMMON_SKILLS:
            # Use word boundaries or literal matching for multi-word skills
            if len(skill) <= 3:
                # Short skills like Go, R, AWS need word boundaries
                pattern = r'\b' + re.escape(skill) + r'\b'
            else:
                pattern = re.escape(skill)
                
            if re.search(pattern, text_lower):
                # Standardize spelling representation
                if skill == "golang":
                    skill = "Go"
                elif skill == "nodejs":
                    skill = "Node.js"
                found_skills.append(skill.title() if len(skill) > 3 else skill.upper())
        
        return list(set(found_skills))

    @staticmethod
    def _extract_section(raw_text: str, section_keywords: List[str]) -> List[str]:
        lines = raw_text.split("\n")
        section_lines = []
        in_section = False
        
        # Section titles are often uppercase or stand-alone lines
        for i, line in enumerate(lines):
            line_str = line.strip()
            if not line_str:
                continue
            
            # Check if this line signals the start of a target section
            line_lower = line_str.lower()
            is_section_header = any(re.search(r'\b' + re.escape(kw) + r'\b', line_lower) for kw in section_keywords)
            
            # Often sections start with a short header (e.g. "Work Experience" or "EDUCATION")
            if is_section_header and len(line_str) < 35:
                in_section = True
                continue
            
            # If we're inside, and we hit another major header keyword that isn't ours, stop
            if in_section:
                other_keywords = ["skills", "experience", "education", "projects", "certifications", "summary", "languages"]
                # Filter out our own keywords
                other_keywords = [ok for ok in other_keywords if ok not in section_keywords]
                is_other_header = any(re.search(r'\b' + re.escape(ok) + r'\b', line_lower) for ok in other_keywords)
                if is_other_header and len(line_str) < 35:
                    in_section = False
                    break
                
                # Add bullet points or non-empty lines
                if len(line_str) > 3:
                    section_lines.append(line_str)
                    if len(section_lines) >= 40:  # Cap length of extracted sections
                        break
                        
        return section_lines if section_lines else []

    @classmethod
    def _extract_job_title(cls, raw_text: str) -> Optional[str]:
        lines = [line.strip() for line in raw_text.split("\n") if line.strip()]
        if len(lines) > 1:
            # Skip the name (first line)
            start_idx = 1
            # Check lines 1 to 5
            for line in lines[start_idx:min(6, len(lines))]:
                line_lower = line.lower()
                # Skip lines containing common contact indicators
                if any(kw in line_lower for kw in ["email", "phone", "contact", "@", "http", "www", "|", "summary", "experience", "education", "skills"]):
                    continue
                # Match typical job title keywords
                title_keywords = [
                    "engineer", "developer", "programmer", "analyst", "scientist", 
                    "manager", "lead", "architect", "consultant", "designer", 
                    "specialist", "expert", "administrator", "officer", "head", 
                    "director", "coordinator", "strategist"
                ]
                if any(kw in line_lower for kw in title_keywords):
                    if len(line) < 60:
                        return line
        return None
